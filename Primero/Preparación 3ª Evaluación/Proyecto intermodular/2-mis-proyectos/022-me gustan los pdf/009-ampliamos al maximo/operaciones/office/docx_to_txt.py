#!/usr/bin/env python3

import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx is not installed. Install it with: sudo apt install python3-docx")
    sys.exit(1)


def ensure_dir(path: Path):
    path.mkdir(parents=True, exist_ok=True)


def collect_docx(inputs):
    files = []
    for item in inputs:
        p = Path(item)
        if p.is_file() and p.suffix.lower() == ".docx":
            files.append(p.resolve())
        elif p.is_dir():
            for child in sorted(p.iterdir()):
                if child.is_file() and child.suffix.lower() == ".docx":
                    files.append(child.resolve())
    return files


def extract_text(docx_path: Path) -> str:
    doc = Document(docx_path)
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))

    return "\n".join(lines)


def main():
    """
    Usage:
        python3 docx_to_txt.py <output_dir> <file_or_folder> [more...]

    Requires:
        sudo apt install python3-docx
    """

    if len(sys.argv) < 3:
        print("Usage:")
        print("  python3 docx_to_txt.py <output_dir> <file_or_folder> [more_files_or_folders...]")
        sys.exit(1)

    output_dir = Path(sys.argv[1]).resolve()
    inputs = sys.argv[2:]

    ensure_dir(output_dir)

    files = collect_docx(inputs)

    if not files:
        print("ERROR: No DOCX files found.")
        sys.exit(1)

    processed = 0
    errors = []

    for idx, f in enumerate(files, start=1):
        try:
            text = extract_text(f)
            out_file = output_dir / f"{idx:03d}_{f.stem}.txt"
            with open(out_file, "w", encoding="utf-8") as out:
                out.write(text)
            processed += 1
        except Exception as e:
            errors.append(f"{f}: {e}")

    report = output_dir / "report.txt"
    with open(report, "w", encoding="utf-8") as r:
        r.write("DOCX TO TXT REPORT\n")
        r.write("==================\n\n")
        r.write(f"Processed: {processed}\n")
        r.write(f"Errors: {len(errors)}\n\n")

        if errors:
            r.write("ERRORS\n------\n")
            for e in errors:
                r.write(e + "\n")

    if processed == 0:
        print("ERROR: No files converted.")
        print(f"REPORT: {report}")
        sys.exit(1)

    print("OK")
    print(f"Processed: {processed}")
    print(f"Errors: {len(errors)}")
    print(f"OUTPUT_DIR: {output_dir}")
    print(f"REPORT: {report}")


if __name__ == "__main__":
    main()
